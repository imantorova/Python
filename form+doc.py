# -*- coding: utf-8 -*-
import sys
from PyQt5 import QtWidgets
from PyQt5 import QtCore
from PyQt5. QtWidgets import QMessageBox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Forma(QtWidgets.QWidget):

    def __init__(self):
        super().__init__()
        self.initUI()


    def initUI(self):
        self.resize(634, 379)
        self.setGeometry(QtCore.QRect(300, 300, 631, 341))
        self.centralwidget = QtWidgets.QWidget(self)
        self.setWindowTitle('Программа для автозаполнения документов')

        self.text_lst = []
        self.text_lst_2 = []
        self.text_lst_3 = []
        self.text_lst_4 = 0
        self.all_form=[]

        #надписи
        self.l_lst={
            'text':["Фамилия","имя","отчество", "место рождения", "паспорт: серия",
                    "номер","кем выдан","адрес регистрации","№ квартиры","этаж","площадь квартиры с балконом",
                    "площадь квартиры", "жилая площадь","площадь кухни","Кадастровый номер","Запись регистрации",
                    "Сумма договора","дата рождения","пол","количество комнат","когда выдан","дата регистрации","руб."],
            'x':[2,210,400,179,2,220,2,2,2,190,2,208,365,500,2,2,2,2,456,337,462,441,440],
            'y':[0,0,0,30,60,60,90,120,150,150,180,180,180,180,210,240,270,30,30,150,60,240,270],
            'width':[44,18,47,83,77,30,54,97,66,25,158,102,81,82,102,102,82,79,18,98,66,92,31],
            'height':[16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16]
            }
        for i in range (23):

            self.label = QtWidgets.QLabel(self)
            self.label.setGeometry(QtCore.QRect(self.l_lst['x'][i], self.l_lst['y'][i], self.l_lst['width'][i], self.l_lst['height'][i]))
            self.label.setText(self.l_lst['text'][i])


        #поля ввода текста
        self.e_lst={
            'x':[53,240,459,270,85,270,84,105,74,221,163,315,455,584,110,110,110],
            'y':[0,0,0,30,60,60,90,120,150,150,180,180,180,180,210,240,270],
            'width':[141,141,161,171,111,171,541,521,110,110,41,41,41,41,321,321,321],
            'height':[20,20,20,20,20,20,20,20,20,20,20,20,20,20,20,20,20]

            }
        self.lE_lst=[]
        for i in range (17):

            self.lineEdit = QtWidgets.QLineEdit(self)
            self.lineEdit.setGeometry(QtCore.QRect(self.e_lst['x'][i], self.e_lst['y'][i], self.e_lst['width'][i], self.e_lst['height'][i]))
            self.lE_lst.append(self.lineEdit)

        #переключатели для типа "пол"
        self.radioButton = QtWidgets.QRadioButton(self)
        self.radioButton.setGeometry(QtCore.QRect(480, 30, 66, 18))
        self.radioButton.setText("мужской")
        self.radioButton.clicked.connect(lambda *args: self.showDialog3(1))

        self.radioButton_2 = QtWidgets.QRadioButton(self)
        self.radioButton_2.setGeometry(QtCore.QRect(552, 30, 66, 18))
        self.radioButton_2.setText("женский")
        self.radioButton_2.clicked.connect(lambda *args: self.showDialog3(2))


        #объединение в группу переключателей
        self.buttonGroup = QtWidgets.QButtonGroup(self)
        self.buttonGroup.addButton(self.radioButton)
        self.buttonGroup.addButton(self.radioButton_2)

        #-------
        #переключатели для типа "количество комнат"
        self.radioButton_3 = QtWidgets.QRadioButton(self)
        self.radioButton_3.setGeometry(QtCore.QRect(441, 150, 29, 18))
        self.radioButton_3.setText("1")
        self.radioButton_3.clicked.connect(lambda *args: self.showDialog4(1))


        self.radioButton_4 = QtWidgets.QRadioButton(self)
        self.radioButton_4.setGeometry(QtCore.QRect(476, 150, 29, 18))
        self.radioButton_4.setText("2")
        self.radioButton_4.clicked.connect(lambda *args: self.showDialog4(2))

        self.radioButton_5 = QtWidgets.QRadioButton(self)
        self.radioButton_5.setGeometry(QtCore.QRect(511, 150, 29, 18))
        self.radioButton_5.setText("3")
        self.radioButton_5.clicked.connect(lambda *args: self.showDialog4(3))

        #объединение в группу переключателей
        self.buttonGroup_2 = QtWidgets.QButtonGroup(self)
        self.buttonGroup_2.addButton(self.radioButton_3)
        self.buttonGroup_2.addButton(self.radioButton_4)
        self.buttonGroup_2.addButton(self.radioButton_5)


        #поле типа дата
        self.dateEdit=QtWidgets.QDateEdit(self)
        self.dateEdit.setGeometry(QtCore.QRect(86, 30, 85, 20))
        self.dateEdit.setDateTime(QtCore.QDateTime(QtCore.QDate(1900, 1, 1)))

        self.dateEdit_2 = QtWidgets.QDateEdit(self)
        self.dateEdit_2.setGeometry(QtCore.QRect(534, 60, 85, 20))
        self.dateEdit_2.setDateTime(QtCore.QDateTime(QtCore.QDate(2000, 1, 1)))

        self.dateEdit_3 = QtWidgets.QDateEdit(self)
        self.dateEdit_3.setGeometry(QtCore.QRect(539, 240, 85, 20))
        self.dateEdit_3.setDateTime(QtCore.QDateTime(QtCore.QDate(2015, 1, 1)))


        self.pushButton = QtWidgets.QPushButton("Проверить",self)
        self.pushButton.setGeometry(QtCore.QRect(430, 310, 75, 23))
        self.pushButton.clicked.connect(self.showDialog)

        self.pushButton_2 = QtWidgets.QPushButton("Сохранить",self)
        self.pushButton_2.setGeometry(QtCore.QRect(550, 310, 75, 23))
        self.pushButton_2.clicked.connect(self.showDialog2)

        self.show()



    def showDialog(self,text):

        self.text_1=self.dateEdit.text()
        self.text_lst_2.append(self.text_1)
        self.text_2=self.dateEdit_2.text()
        self.text_lst_2.append(self.text_2)
        self.text_3=self.dateEdit_3.text()
        self.text_lst_2.append(self.text_3)


        for i in range (17):
            self.text=self.lE_lst[i].text()
            self.text_lst.append(self.text)

        for i in (0,1,2):
            if self.text_lst[i].isalpha()==False:
                self.lE_lst[i].setText('')
                self.reply = QtWidgets.QMessageBox.question(self, 'Ошибка!'
                                                                  'Поле введено неверно! '
                                                                  'Должны быть только буквы!',self.l_lst['text'][i], QMessageBox.Yes)

        for i in (4,5,8,9,16):
            if self.text_lst[i].isdigit()==False:
                self.lE_lst[i].setText('')
                self.reply2 = QtWidgets.QMessageBox.question(self, 'Ошибка!'
                                                                  'Поле введено неверно! '
                                                                  'Должны быть только цифры!',self.l_lst['text'][i], QMessageBox.Yes)


    def showDialog3(self,i):

        if i==1:
            self.text_lst_3.append('мужской')
            self.all_form = self.text_lst_3
        elif i==2:
            self.text_lst_3.append('женский')
            self.all_form = self.text_lst_3

    def showDialog4(self,i):

        if i==1:
            self.text_lst_4 = 1
            self.all_form.append(self.text_lst_4)
        elif i==2:
            self.text_lst_4 = 2
            self.all_form.append(self.text_lst_4)
        elif i==3:
            self.text_lst_4 = 3
            self.all_form.append(self.text_lst_4)


    def showDialog2(self):

        self.all_form = self.all_form+self.text_lst+self.text_lst_2
        document = Document()
        p = document.add_paragraph('ДОГОВОР купли – продажи квартиры № %(10)s '''% {"10":self.all_form[10]})
        p = document.add_paragraph('г. Санкт-Петербург                    	   «__» ____________ 201__ года')
        p = document.add_paragraph('Общество с ограниченной ответственностью «Карат», являющееся юридическим лицом по законодательству'
                                        ' Российской Федерации, зарегистрированное в МИ МНС № 11 по Санкт-Петербургу Свидетельством '
                                        'о государственной регистрации ЮЛ серия 78 № 004629954 от 31.10.2003 года, ОГРН 1037869011421, '
                                        'ИНН 7842003710, КПП 780601001, местонахождение: 195112, г. Санкт-Петербург, Малоохтинский пр., д. 61,'
                                        ' литера А, пом. 61, в лице генерального директора Осипова Дмитрия Вячеславовича, действующего на'
                                        ' основании Устава, именуемое далее «Продавец», и')

        p = document.add_paragraph('Гражданин Российской Федерации %(2)s %(3)s %(4)s %(19)s года рождения, место рождения:%(5)s, пол: %(0)s, паспорт %(6)s %(7)s выдан %(20)s г. '
                                        '%(8)s, зарегистрированный по адресу (адрес для уведомлений): %(9)s, именуемый далее «Покупатель», с другой стороны, '
                                        'заключили настоящий Договор о нижеследующем'% {"0":self.all_form[0], "2":self.all_form[2],"3":self.all_form[3],
                                                                                        "4":self.all_form[4],"5":self.all_form[5],"6":self.all_form[6],
                                                                                        "7":self.all_form[7],"8":self.all_form[8],"9":self.all_form[9],
                                                                                        "19":self.all_form[19],"20":self.all_form[20]})
        p = document.add_paragraph('Продавец продал, а Покупатель купил %(1)sкомнатную квартиру, находящуюся по адресу: гор. Санкт-Петербург, г. Сестрорецк,'
                                   'Приморское шоссе, д. 293, кв. %(10)s'%{"1":self.all_form[1],"10":self.all_form[10]},style='ListNumber')
        
        p = document.add_paragraph('Указанная квартира (кадастровый номер: %(16)s) расположена на %(11)s этаже жилом 10-17-ти этажном доме'
                                   'со встроенными помещениями 2014 года постройки. Общая площадь квартиры составляет – %(13)s кв.м.;'
                                   'из нее жилая площадь – %(14)s кв.м.'%{"16":self.all_form[16], "11":self.all_form[11],"13":self.all_form[13],
                                                                          "14":self.all_form[14]},style='ListNumber')
        
        p = document.add_paragraph('Отчуждаемая квартира принадлежит Обществу с ограниченной ответственностью «Карат» на праве собственности,'
                                   'о чем в Едином государственном реестре прав на недвижимое имущество и сделок с ним'
                                   '%(21)s года сделана запись регистрации № %(17)s.'% {"17":self.all_form[17], "21":self.all_form[21]},style='ListNumber')
        
        p = document.add_paragraph('Указанная квартира по договоренности сторон продается и покупается за сумму %(18)s рублей 00 копеек. Взаиморасчеты произведены'
                                   'в полном объеме между Сторонами до подписания настоящего Договора. Стороны финансовых и иных претензий друг к другу не имеют.'%
                                   {"18":self.all_form[18]},style='ListNumber')
        
        p = document.add_page_break()

        p = document.add_paragraph('АКТ ПРИЕМА-ПЕРЕДАЧИ КВАРТИРЫ  № %(10)s '''% {"10":self.all_form[10]})
        p = document.add_paragraph('г. Санкт-Петербург                    	   «__» ____________ 201__ года')
        p = document.add_paragraph('Общество с ограниченной ответственностью «Карат», являющееся юридическим лицом по законодательству'
                                        ' Российской Федерации, зарегистрированное в МИ МНС № 11 по Санкт-Петербургу Свидетельством '
                                        'о государственной регистрации ЮЛ серия 78 № 004629954 от 31.10.2003 года, ОГРН 1037869011421, '
                                        'ИНН 7842003710, КПП 780601001, местонахождение: 195112, г. Санкт-Петербург, Малоохтинский пр., д. 61,'
                                        ' литера А, пом. 61, в лице генерального директора Осипова Дмитрия Вячеславовича, действующего на'
                                        ' основании Устава, именуемое далее «Продавец», и')

        p = document.add_paragraph('Гражданин Российской Федерации %(2)s %(3)s %(4)s %(19)s года рождения, место рождения:%(5)s, пол: %(0)s, паспорт %(6)s %(7)s выдан %(20)s г. '
                                        '%(8)s, зарегистрированный по адресу (адрес для уведомлений): %(9)s, именуемый далее «Покупатель», с другой стороны, '
                                        'другой стороны, составили настоящий акт приема-передачи квартиры № %(10)s в соответствии с договором купли-продажи '
                                        'квартиры № %(10)s от «__» ______ 201__ года.'% {"0":self.all_form[0], "2":self.all_form[2],"3":self.all_form[3],
                                                                                        "4":self.all_form[4],"5":self.all_form[5],"6":self.all_form[6],
                                                                                        "7":self.all_form[7],"8":self.all_form[8],"9":self.all_form[9],
                                                                                        "10":self.all_form[10],"19":self.all_form[19],"20":self.all_form[20]})
        
        p = document.add_paragraph('Квартира передается в том виде и состоянии, в котором она находится на момент передачи ее Покупателю. С момента подписания'
                                   'настоящего акта сторонами Продавец передает Покупателю, а Покупатель принимает квартиру, расположенную по адресу: '
                                   'г. Санкт-Петербург, г. Сестрорецк, Приморское шоссе, д. 293, следующих характеристик:',style='ListNumber')
        
        table = document.add_table(rows=2, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Номер квартиры'
        hdr_cells[1].text = 'Количество комнат'
        hdr_cells[2].text = 'Этаж'
        hdr_cells[3].text = 'Общая жилая площадь по данным ПИБ, кв.м.'
        hdr_cells[4].text = 'Общая площадь с учетом приведенной площади балкона, лоджии, кв.м.'
        hdr_cells[5].text = 'Общая внутренняя площадь квартиры, кв.м.'

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = ' '
        hdr_cells[1].text = ' '
        hdr_cells[2].text = ' '
        hdr_cells[3].text = ' '
        hdr_cells[4].text = ' '
        hdr_cells[5].text = ' '
        
        document.save('demo.docx')

        self.reply3 = QtWidgets.QMessageBox.question(self, 'Сохранение документа','Документ успешно сохранен!', QMessageBox.Yes)
        
if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    forma = Forma()
    sys.exit(app.exec_())
