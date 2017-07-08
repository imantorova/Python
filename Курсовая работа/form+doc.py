'''
Программа для автозаполнения форм договоров и актов по жилым помещениям (квартирам) по объекту: г. Сестрорецк,
Приморское шоссе, д. 293 на основании введенных пользователем данных (техническая документация по объекту и
личные данные покупателей объекта).
'''

# -*- coding: utf-8 -*-

import sys

from PyQt5 import QtWidgets
from PyQt5 import QtCore
from PyQt5. QtWidgets import QMessageBox

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Forma(QtWidgets.QWidget):

    def __init__(self):
        super().__init__()
        self.initUI()


    def initUI(self):

        #фоновый виджет
        self.resize(634, 379)
        self.setGeometry(QtCore.QRect(300, 300, 631, 341))
        self.centralwidget = QtWidgets.QWidget(self)
        self.setWindowTitle('Программа заполнения документов для объекта: г. Сестрорецк, Приморское шоссе, д. 293')

        self.text_lst = []
        self.text_lst_2 = []
        self.text_lst_3 = 0
        self.text_lst_4 = 0
        self.all_form = ['#Значение!#','#Значение!#']

        #надписи
        self.l_lst = {
            'text':["Фамилия","имя","отчество", "место рождения", "паспорт: серия",
                    "номер","кем выдан","адрес регистрации","№ квартиры","этаж","площадь квартиры с балконом",
                    "площадь квартиры", "жилая площадь","площадь кухни","Кадастровый номер","Запись регистрации",
                    "Сумма договора","дата рождения","пол","количество комнат","когда выдан","дата регистрации","руб."],
            'x':[2,210,400,179,2,220,2,2,2,190,2,208,365,500,2,2,2,2,456,337,462,441,440],
            'y':[0,0,0,30,60,60,90,120,150,150,180,180,180,180,210,240,270,30,30,150,60,240,270],
            'width':[44,18,47,83,77,30,54,97,66,25,158,102,81,82,102,102,82,79,18,98,66,92,31],
            'height':[16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16,16]
            }

        for i in range (23):
            self.label = QtWidgets.QLabel(self)
            self.label.setGeometry(QtCore.QRect(self.l_lst['x'][i], self.l_lst['y'][i], self.l_lst['width'][i], self.l_lst['height'][i]))
            self.label.setText(self.l_lst['text'][i])


        #поля ввода текста
        self.e_lst = {
            'x':[53,240,459,270,85,270,84,105,74,221,163,315,455,584,110,110,110],
            'y':[0,0,0,30,60,60,90,120,150,150,180,180,180,180,210,240,270],
            'width':[141,141,161,171,111,171,541,521,110,110,41,41,41,41,321,321,321],
            'height':[20,20,20,20,20,20,20,20,20,20,20,20,20,20,20,20,20]
            }

        self.lE_lst = []

        for i in range (17):
            self.lineEdit = QtWidgets.QLineEdit(self)
            self.lineEdit.setGeometry(QtCore.QRect(self.e_lst['x'][i], self.e_lst['y'][i], self.e_lst['width'][i], self.e_lst['height'][i]))

            #для полей данных "Фамилия", "имя", "отчество", "место рождения" назначается маска ввода "только буквы"
            if i == 0:
                self.lineEdit.setInputMask ("AAAAAAAAAAAAAAAAAAAA")

            if i == 1:
                self.lineEdit.setInputMask ("AAAAAAAAAAAAAAAAAAAA")

            if i == 2:
                self.lineEdit.setInputMask ("AAAAAAAAAAAAAAAAAAAA")

            if i == 3:
                self.lineEdit.setInputMask ("AAAAAAAAAAAAAAAAAAAA")

            '''
           для полей данных "паспорт: серия", "номер", "№ квартиры", "этаж", "площадь квартиры с балконом",
           "площадь квартиры", "жилая площадь","площадь кухни", "Сумма договора" назначается маска ввода 
           "только цифры" с соответствующим количеством символов
           '''
            if i == 4:
                self.lineEdit.setInputMask ("99 99")

            if i == 5:
                self.lineEdit.setInputMask ("999999")

            if i == 8:
                self.lineEdit.setInputMask ("999")

            if i == 9:
                self.lineEdit.setInputMask ("99")

            if i == 10:
                self.lineEdit.setInputMask ("99,99")

            if i == 11:
                self.lineEdit.setInputMask ("99,99")

            if i == 12:
                self.lineEdit.setInputMask ("99,99")

            if i == 13:
                self.lineEdit.setInputMask ("99,99")

            if i == 16:
                self.lineEdit.setInputMask ("9 999 999")

            #параметры всех полей ввода текста записаны в список
            self.lE_lst.append(self.lineEdit)


        #переключатели для типа "пол"
        self.radioButton = QtWidgets.QRadioButton(self)
        self.radioButton.setGeometry(QtCore.QRect(480, 30, 66, 18))
        self.radioButton.setText("мужской")
        self.radioButton.clicked.connect(lambda *args: self.Pol_str(1))

        self.radioButton_2 = QtWidgets.QRadioButton(self)
        self.radioButton_2.setGeometry(QtCore.QRect(552, 30, 66, 18))
        self.radioButton_2.setText("женский")
        self.radioButton_2.clicked.connect(lambda *args: self.Pol_str(2))


        #объединение в группу переключателей
        self.buttonGroup = QtWidgets.QButtonGroup(self)
        self.buttonGroup.addButton(self.radioButton)
        self.buttonGroup.addButton(self.radioButton_2)

        #-------
        #переключатели для типа "количество комнат"
        self.radioButton_3 = QtWidgets.QRadioButton(self)
        self.radioButton_3.setGeometry(QtCore.QRect(441, 150, 29, 18))
        self.radioButton_3.setText("1")
        self.radioButton_3.clicked.connect(lambda *args: self.Komnaty_str(1))


        self.radioButton_4 = QtWidgets.QRadioButton(self)
        self.radioButton_4.setGeometry(QtCore.QRect(476, 150, 29, 18))
        self.radioButton_4.setText("2")
        self.radioButton_4.clicked.connect(lambda *args: self.Komnaty_str(2))

        self.radioButton_5 = QtWidgets.QRadioButton(self)
        self.radioButton_5.setGeometry(QtCore.QRect(511, 150, 29, 18))
        self.radioButton_5.setText("3")
        self.radioButton_5.clicked.connect(lambda *args: self.Komnaty_str(3))

        #объединение в группу переключателей
        self.buttonGroup_2 = QtWidgets.QButtonGroup(self)
        self.buttonGroup_2.addButton(self.radioButton_3)
        self.buttonGroup_2.addButton(self.radioButton_4)
        self.buttonGroup_2.addButton(self.radioButton_5)


        #поле типа "дата"
        self.dateEdit=QtWidgets.QDateEdit(self)
        self.dateEdit.setGeometry(QtCore.QRect(86, 30, 85, 20))
        self.dateEdit.setDateTime(QtCore.QDateTime(QtCore.QDate(1900, 1, 1)))

        self.dateEdit_2 = QtWidgets.QDateEdit(self)
        self.dateEdit_2.setGeometry(QtCore.QRect(534, 60, 85, 20))
        self.dateEdit_2.setDateTime(QtCore.QDateTime(QtCore.QDate(2000, 1, 1)))

        self.dateEdit_3 = QtWidgets.QDateEdit(self)
        self.dateEdit_3.setGeometry(QtCore.QRect(539, 240, 85, 20))
        self.dateEdit_3.setDateTime(QtCore.QDateTime(QtCore.QDate(2015, 1, 1)))

        #кнопка для преобразования данных в строки и запись данных в файл .docx
        self.pushButton = QtWidgets.QPushButton("Сохранить",self)
        self.pushButton.setGeometry(QtCore.QRect(550, 310, 75, 23))
        self.pushButton.clicked.connect(self.Save_data)

        self.show()

    def Pol_str(self,i):

        #считывание выбранного варианта для данных типа "пол"
        if i == 1:
            self.text_lst_3 = 'мужской'
            self.all_form[0] = self.text_lst_3

        elif i == 2:
            self.text_lst_3 = 'женский'
            self.all_form[0] = self.text_lst_3

    def Komnaty_str(self,i):

        #считывание выбранного варианта для данных типа "количество комнат"
        if i == 1:
            self.text_lst_4 = 1
            self.all_form[1] = self.text_lst_4

        elif i == 2:
            self.text_lst_4 = 2
            self.all_form[1] = self.text_lst_4

        elif i == 3:
            self.text_lst_4 = 3
            self.all_form[1] = self.text_lst_4

    def Save_data(self):

        #считывание введенных данных из текстовых полей
        for i in range (17):
            self.text = self.lE_lst[i].text()
            if self.text == '':
                self.text = '#Значение!#'
            self.text_lst.append(self.text)

        #считывание выбранного варианта для данных типа "дата"
        self.text_1 = self.dateEdit.text()
        self.text_lst_2.append(self.text_1)
        self.text_2 = self.dateEdit_2.text()
        self.text_lst_2.append(self.text_2)
        self.text_3 = self.dateEdit_3.text()
        self.text_lst_2.append(self.text_3)

        #формирование списка из считанных данных для последующего переноса в файл .docx
        self.all_form = self.all_form+self.text_lst+self.text_lst_2

        #заполнение файла .docx(первая часть - договор)
        document = Document()

        p = document.add_paragraph('Договор купли – продажи квартиры № %(10)s '% {"10":self.all_form[10]})
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('г. Санкт-Петербург')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p = document.add_paragraph('«__» ____________ 201__ года')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = document.add_paragraph('Общество с ограниченной ответственностью «Карат», являющееся юридическим лицом по законодательству'
                                        ' Российской Федерации, зарегистрированное в МИ МНС № 11 по Санкт-Петербургу Свидетельством '
                                        'о государственной регистрации ЮЛ серия 78 № 004629954 от 31.10.2003 года, ОГРН 1037869011421, '
                                        'ИНН 7842003710, КПП 780601001, местонахождение: 195112, г. Санкт-Петербург, Малоохтинский пр., д. 61,'
                                        ' литера А, пом. 61, в лице генерального директора Осипова Дмитрия Вячеславовича, действующего на'
                                        ' основании Устава, именуемое далее «Продавец», и')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Гражданин Российской Федерации %(2)s %(3)s %(4)s %(19)s года рождения, место рождения:%(5)s, пол: %(0)s, '
                                   'паспорт %(6)s %(7)s выдан %(20)s г. %(8)s, зарегистрированный по адресу (адрес для уведомлений): %(9)s, '
                                   'именуемый далее «Покупатель», с другой стороны,'
                                   ' заключили настоящий Договор о нижеследующем:'%{"0":self.all_form[0],"2":self.all_form[2],"3":self.all_form[3],
                                    "4":self.all_form[4],"5":self.all_form[5],"6":self.all_form[6],
                                    "7":self.all_form[7],"8":self.all_form[8],"9":self.all_form[9],
                                    "19":self.all_form[19],"20":self.all_form[20]})
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Продавец продал, а Покупатель купил %(1)s-комнатную квартиру, находящуюся по адресу: гор. Санкт-Петербург,'
                                   ' г. Сестрорецк, Приморское шоссе, д. 293, кв. %(10)s.'%{"1":self.all_form[1],"10":self.all_form[10]},
                                   style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Указанная квартира (кадастровый номер: %(16)s) расположена на %(11)s этаже жилом 10-17-ти этажном доме'
                                   ' со встроенными помещениями 2014 года постройки. Общая площадь квартиры составляет – %(13)s кв.м.;'
                                   ' из нее жилая площадь – %(14)s кв.м.'%{"16":self.all_form[16], "11":self.all_form[11],"13":self.all_form[13],
                                    "14":self.all_form[14]},style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Отчуждаемая квартира принадлежит Обществу с ограниченной ответственностью «Карат» на праве собственности,'
                                   'о чем в Едином государственном реестре прав на недвижимое имущество и сделок с ним '
                                   '%(21)s года сделана запись регистрации № %(17)s.'%{"17":self.all_form[17], "21":self.all_form[21]},
                                   style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Указанная квартира по договоренности сторон продается и покупается за сумму     %(18)s рублей 00 копеек. '
                                   'Взаиморасчеты произведены в полном объеме между Сторонами до подписания настоящего Договора. Стороны '
                                   'финансовых и иных претензий друг к другу не имеют.'% {"18":self.all_form[18]},style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Продавец гарантирует, что до заключения настоящего Договора отчуждаемая квартира никому не продана, не'
                                   ' подарена, не заложена, в споре, под арестом (запрещением) не состоит, правами третьих лиц не обременена.'
                                   ,style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Продавец гарантирует, что на момент заключения настоящего Договора в отчуждаемой квартире нет '
                                   'зарегистрированных лиц.',style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Продавец передает Покупателю квартиру в течение трех дней с момента подписания настоящего Договора. '
                                   'Квартира передается в том виде и состоянии, в котором она находится на момент передачи ее Покупателю. '
                                   'Покупатель не имеет претензий к Продавцу по срокам передачи вышеуказанной квартиры по акту приема-передачи,'
                                   ' по техническому состоянию передаваемой квартиры, по размеру площади передаваемой квартиры, претензий '
                                   'по оборудованию и состоянию мест общего пользования. Покупатель несет все имущественные риски, связанные с '
                                   'гибелью или порчей имущества (квартиры) и общего имущества Объекта, а также все расходы по их содержанию с '
                                   'даты подписания акта приема-передачи. Право собственности на квартиру у Покупателя возникает с момента '
                                   'государственной регистрации перехода права собственности на квартиру № %(10)s.'% {"10":self.all_form[10]}
                                   ,style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Покупатель не имеет права до полного оформления квартиры в собственность  (т.е. получения Свидетельства о'
                                   ' праве собственности на квартиру) предпринимать какие-либо действия по перепланировке квартиры, ее '
                                   'инженерного обеспечения, по электроснабжению, теплоснабжению, холодному и горячему водоснабжению, '
                                   'канализации. В случае нарушения данного пункта Договора вся ответственность возлагается на Покупателя.'
                                   ,style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Покупатель за свой счет осуществляет эксплуатацию и ремонт квартиры, а также участвует в расходах,'
                                   ' связанных с техническим обслуживанием и ремонтом, в том числе и капитальным, всего дома соразмерно'
                                   ' с занимаемой площадью.',style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Настоящий Договор составлен и подписан в трех экземплярах, имеющих одинаковую юридическую силу.'
                                   ,style='ListNumber')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.fio = self.all_form[2]+' '+self.all_form[3]+' '+self.all_form[4]
        self.fio_2 = '__________________/'+self.all_form[2]+' '+self.all_form[3][0]+'. '+self.all_form[4][0]+'./'


        table = document.add_table(rows=3, cols=2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Продавец'
        hdr_cells[1].text = 'Покупатель'

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'Генеральный директор ООО "Карат"'
        hdr_cells[1].text = (self.fio)

        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = '__________________/Осипов Д. В./'
        hdr_cells[1].text = (self.fio_2)

        table.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        #разрыв страницы
        #заполнение файла .docx(вторая часть - акт)
        p = document.add_page_break()

        p = document.add_paragraph('АКТ ПРИЕМА-ПЕРЕДАЧИ КВАРТИРЫ  № %(10)s '% {"10":self.all_form[10]})
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('г. Санкт-Петербург')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p = document.add_paragraph('«__» ____________ 201__ года')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = document.add_paragraph('Общество с ограниченной ответственностью «Карат», являющееся юридическим лицом по законодательству'
                                        ' Российской Федерации, зарегистрированное в МИ МНС № 11 по Санкт-Петербургу Свидетельством '
                                        'о государственной регистрации ЮЛ серия 78 № 004629954 от 31.10.2003 года, ОГРН 1037869011421, '
                                        'ИНН 7842003710, КПП 780601001, местонахождение: 195112, г. Санкт-Петербург, Малоохтинский пр., д. 61,'
                                        ' литера А, пом. 61, в лице генерального директора Осипова Дмитрия Вячеславовича, действующего на'
                                        ' основании Устава, именуемое далее «Продавец», и')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('Гражданин Российской Федерации %(2)s %(3)s %(4)s %(19)s года рождения, место рождения:%(5)s, пол: %(0)s, '
                                   'паспорт %(6)s %(7)s выдан %(20)s г. %(8)s, зарегистрированный по адресу (адрес для уведомлений): %(9)s, '
                                   'именуемый далее «Покупатель», с другой стороны, другой стороны, составили настоящий акт приема-передачи'
                                   ' квартиры № %(10)s в соответствии с договором купли-продажи квартиры № %(10)s от «__» ______ 201__ '
                                   'года.'%{"0":self.all_form[0],"2":self.all_form[2],"3":self.all_form[3],
                                            "4":self.all_form[4],"5":self.all_form[5],"6":self.all_form[6],
                                            "7":self.all_form[7],"8":self.all_form[8],"9":self.all_form[9],
                                            "10":self.all_form[10],"19":self.all_form[19],"20":self.all_form[20]})
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('1. Квартира передается в том виде и состоянии, в котором она находится на момент передачи ее Покупателю.'
                                   ' С момента подписания настоящего акта сторонами Продавец передает Покупателю, а Покупатель принимает квартиру,'
                                   ' расположенную по адресу: г. Санкт-Петербург, г. Сестрорецк, Приморское шоссе, д. 293, следующих характеристик:')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        table = document.add_table(rows=2, cols=7)
        table.style = 'TableGrid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Номер квартиры'
        hdr_cells[1].text = 'Количество комнат'
        hdr_cells[2].text = 'Этаж'
        hdr_cells[3].text = 'Общая жилая площадь по данным ПИБ, кв.м.'
        hdr_cells[4].text = 'Кухня по данным ПИБ, кв.м.'
        hdr_cells[5].text = 'Общая площадь с учетом приведенной площади балкона, лоджии, кв.м.'
        hdr_cells[6].text = 'Общая внутренняя площадь квартиры, кв.м.'

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = str(self.all_form[10])
        hdr_cells[1].text = str(self.all_form[1])
        hdr_cells[2].text = str(self.all_form[11])
        hdr_cells[3].text = str(self.all_form[14])
        hdr_cells[4].text = str(self.all_form[15])
        hdr_cells[5].text = str(self.all_form[12])
        hdr_cells[6].text = str(self.all_form[13])


        p = document.add_paragraph('2. Покупатель не имеет к Продавцу претензий по качеству вышеуказанной квартиры № %(10)s. Покупатель '
                                   'не имеет претензий к Продавцу по срокам передачи вышеуказанной квартиры по настоящему акту, по техническому'
                                   ' состоянию передаваемой квартиры, по размеру площади передаваемой квартиры, претензий по оборудованию и '
                                   'состоянию мест общего пользования.'%{"10":self.all_form[10]})
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('3. Взаиморасчеты между Покупателем и Продавцом произведены полностью.  Окончательная стоимость квартиры '
                                   'составляет %(18)s рублей 00 копеек. Стороны финансовых и иных претензий друг к другу не '
                                   'имеют.'%{"18":self.all_form[18]})
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('4. С момента подписания Покупателем акта приема-передачи Продавец не несет ответственности за сохранность'
                                   ' передаваемой квартиры № %(10)s.'%{"10":self.all_form[10]})
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('5. Покупатель за свой счет осуществляет эксплуатацию и ремонт квартиры, а также участвует в расходах, '
                                   'связанных с техническим обслуживанием и ремонтом, в том числе и капитальным, всего дома соразмерно с занимаемой'
                                   ' площадью.')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('6. Настоящий акт составлен в трех экземплярах, имеющих одинаковую юридическую силу.')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = document.add_paragraph('7. Подписи сторон:')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        table = document.add_table(rows=3, cols=2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Продавец'
        hdr_cells[1].text = 'Покупатель'

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'Генеральный директор ООО "Карат"'
        hdr_cells[1].text = (self.fio)

        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = '__________________/Осипов Д. В./'
        hdr_cells[1].text = (self.fio_2)

        table.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        document.save('Договор+акт кв.№ %(10)s.docx'%{"10":self.all_form[10]})

        self.reply3 = QtWidgets.QMessageBox.question(self, 'Сохранение документа','Документ успешно сохранен!', QMessageBox.Yes)

        sys.exit(0)
if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    forma = Forma()
    sys.exit(app.exec_())
